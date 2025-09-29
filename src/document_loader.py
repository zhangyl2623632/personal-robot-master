import logging
import os
import re
import yaml
import time
import hashlib
import uuid
from typing import List, Dict, Any, Optional, Tuple
from langchain_community.document_loaders import (
    TextLoader,
    DirectoryLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredExcelLoader,
    CSVLoader,
    ImageCaptionLoader
)
from langchain_community.document_loaders.pdf import PyPDFLoader as PDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from PIL import Image
import pytesseract
import langdetect
import numpy as np
from concurrent.futures import ThreadPoolExecutor, as_completed
from src.config import global_config
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# 导入文档分类器
from src.document_classifier import document_classifier
from src.llm_client import extract_metadata_from_text

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentLoader:
    """文档加载器，用于加载和处理不同类型的本地文档"""
    
    def __init__(self, config=None):
        """初始化文档加载器"""
        self.config = config or global_config
        self.processing_config = self._load_processing_config()
        
        # 初始化缓存
        self._cache = {}
        self._cache_enabled = self.processing_config.get('caching', {}).get('enabled', True)
        self._cache_ttl = self.processing_config.get('caching', {}).get('ttl_seconds', 3600)
        
        # 初始化分块策略
        self._init_text_splitters()
        
        # 初始化并发处理池
        concurrency_config = self.processing_config.get('concurrency', {})
        self._executor = ThreadPoolExecutor(
            max_workers=concurrency_config.get('max_workers', 4)
        ) if concurrency_config.get('enabled', True) else None
        
        # 初始化嵌入模型（用于语义分块）
        self._embedding_model = self._init_embedding_model()
    
    def _load_processing_config(self) -> Dict[str, Any]:
        """加载文档处理配置"""
        config_path = os.path.join('config', 'document_processing.yaml')
        try:
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                logger.info(f"成功加载文档处理配置: {config_path}")
                return config
            else:
                logger.warning(f"配置文件不存在: {config_path}，使用默认配置")
                # 返回默认配置
                return self._get_default_processing_config()
        except Exception as e:
            logger.error(f"加载配置文件失败: {str(e)}")
            return self._get_default_processing_config()
    
    def _get_default_processing_config(self) -> Dict[str, Any]:
        """获取默认处理配置"""
        return {
            'chunking_strategies': {
                'default': {
                    'chunk_size': 500,
                    'chunk_overlap': 50,
                    'separators': ['\n\n', '\n', '。', '！', '？', '；', '，', ' ', '']
                }
            },
            'semantic_chunking': {'enabled': True},
            'metadata_extraction': {'enabled': True},
            'ocr': {'enabled': True, 'languages': ['chi_sim', 'eng']}
        }
    
    def _init_text_splitters(self):
        """初始化文本分块器"""
        strategies = self.processing_config.get('chunking_strategies', {})
        
        # 默认分块器
        default_strategy = strategies.get('default', {})
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=default_strategy.get('chunk_size', 500),
            chunk_overlap=default_strategy.get('chunk_overlap', 50),
            separators=default_strategy.get('separators', ['\n\n', '\n', '。', '！', '？', '；', '，', ' ', ''])
        )
        
        # PDF分块器
        pdf_strategy = strategies.get('academic_paper', {}) or default_strategy
        self.pdf_text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=pdf_strategy.get('chunk_size', 800),
            chunk_overlap=pdf_strategy.get('chunk_overlap', 100),
            separators=pdf_strategy.get('separators', ['\n\n', '\n', '。', '！', '？', '；', '\n\t', '', ' '])
        )
        
        # 合同文档分块器
        contract_strategy = strategies.get('contract', {}) or default_strategy
        self.contract_text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=contract_strategy.get('chunk_size', 600),
            chunk_overlap=contract_strategy.get('chunk_overlap', 80),
            separators=contract_strategy.get('separators', ['\n\n', '\n', '。', '！', '？', '；', '条款', '第.*条', '', ' '])
        )
        
        # 技术文档分块器
        tech_strategy = strategies.get('technical_doc', {}) or default_strategy
        self.tech_text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=tech_strategy.get('chunk_size', 700),
            chunk_overlap=tech_strategy.get('chunk_overlap', 90),
            separators=tech_strategy.get('separators', ['\n\n', '\n', '。', '！', '？', '；', '\t', '', ' '])
        )
    
    def _init_embedding_model(self):
        """初始化嵌入模型用于语义分块"""
        if not self.processing_config.get('semantic_chunking', {}).get('enabled', True):
            return None
        
        try:
            from sentence_transformers import SentenceTransformer
            model = SentenceTransformer('BAAI/bge-large-zh-v1.5')
            logger.info("成功初始化嵌入模型用于语义分块")
            return model
        except Exception as e:
            logger.warning(f"初始化嵌入模型失败: {str(e)}，语义分块功能将被禁用")
            return None
    
    def load_document(self, file_path: str) -> List[Document]:
        """加载单个文档，支持语义分块和增强元数据提取，与_get_loader_cls方法兼容"""
        logger.info(f"开始加载文档: {file_path}")
        
        # 检查缓存
        cache_key = self._get_cache_key(file_path)
        if self._cache_enabled and cache_key in self._cache:
            cached_data = self._cache[cache_key]
            if time.time() < cached_data['expires']:
                logger.info(f"从缓存加载文档: {file_path}")
                return cached_data['documents']
            else:
                del self._cache[cache_key]  # 缓存过期，删除
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.warning(f"文件不存在: {file_path}")
            return []
            
        # 检查文件大小，避免处理过大的文件
        try:
            max_size_mb = self.processing_config.get('document_validation', {}).get('max_size_mb', 100)
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > max_size_mb:
                logger.warning(f"文件过大: {file_path}, 大小: {file_size_mb:.2f}MB, 限制: {max_size_mb}MB")
                return []
        except Exception as e:
            logger.error(f"获取文件大小失败: {file_path}, 错误: {str(e)}")
            return []
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # 检查文件是否为空
            if os.path.getsize(file_path) == 0:
                logger.warning(f"空文件，跳过加载: {file_path}")
                return []
                
            # 文档验证
            if not self._validate_document(file_path, file_extension):
                return []
                
            # 初始化文档列表
            documents = []
                
            # 特殊处理图片文件
            if file_extension in [".jpg", ".jpeg", ".png", ".gif"]:
                # 调用专门的图片处理方法
                docs = self._process_image_file(file_path)
                # 添加缓存
                if self._cache_enabled:
                    self._cache[cache_key] = {
                        'documents': docs,
                        'expires': time.time() + self._cache_ttl
                    }
                return docs
            else:
                # 使用_get_loader_cls方法获取加载器
                loader_cls = self._get_loader_cls(file_extension)
                if not loader_cls:
                    return []
                
                # 根据文件类型配置加载器
                if file_extension in [".txt", ".md", ".csv"]:
                    loader = loader_cls(file_path, encoding="utf-8")
                else:
                    loader = loader_cls(file_path)
                
                # 加载文档
                documents = loader.load()
            
            # 文档类型分类
            doc_type = self._classify_document_type(documents[0]) if documents else 'unknown'
            
            # 选择分块策略
            split_docs = self._split_documents(documents, file_extension, doc_type)
            
            # 增强元数据
            enhanced_docs = self._enhance_metadata(split_docs, file_path, file_extension, doc_type)
            
            # 添加缓存
            if self._cache_enabled:
                self._cache[cache_key] = {
                    'documents': enhanced_docs,
                    'expires': time.time() + self._cache_ttl
                }
            
            # 如果是临时文件，尝试获取原始文件名
            if "temp" in file_path.lower():
                original_filename = documents[0].metadata.get('source', file_path) if documents else file_path
                logger.info(f"成功加载文档: {file_path} (原始文件名: {original_filename})，生成 {len(enhanced_docs)} 个分块")
            else:
                logger.info(f"成功加载文档: {file_path}，生成 {len(enhanced_docs)} 个分块")
            
            return enhanced_docs
        except UnicodeDecodeError as e:
            logger.error(f"文件编码错误: {str(e)}，文件: {file_path}")
            # 尝试使用OCR处理（如果是图片或扫描PDF）
            if file_extension in ['.pdf', '.jpg', '.jpeg', '.png']:
                return self._extract_text_with_ocr(file_path)
            return []
        except Exception as e:
            logger.error(f"加载文档失败: {file_path}, 错误: {str(e)}")
            # 对于图片文件，尝试OCR
            if file_extension in ['.jpg', '.jpeg', '.png']:
                return self._extract_text_with_ocr(file_path)
            return []
    
    def _process_image_file(self, file_path: str) -> List[Document]:
        """处理图片文件，使用OCR提取文本，优化图片处理逻辑"""
        try:
            # 获取图片处理配置
            ocr_config = self.processing_config.get('ocr', {})
            languages = '+'.join(ocr_config.get('languages', ['chi_sim', 'eng']))
            
            logger.info(f"使用OCR处理图片文件: {file_path}")
            
            # 打开图片文件
            with Image.open(file_path) as img:
                # 预处理图片以提高OCR质量
                processed_img = self._preprocess_image(img)
                
                # 使用OCR提取文本
                text = pytesseract.image_to_string(processed_img, lang=languages)
                
                if not text.strip():
                    logger.warning(f"OCR未能从图片提取到文本: {file_path}")
                    return []
                
                # 检测语言
                language = self._detect_language(text)
                
                # 创建文档对象
                document = Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "file_name": os.path.basename(file_path),
                        "file_type": file_extension,
                        "language": language,
                        "ocr_processed": True,
                        "ocr_languages": languages,
                        "processed_at": time.strftime('%Y-%m-%d %H:%M:%S')
                    }
                )
                
                # 分块处理OCR提取的文本
                from langchain_text_splitters import RecursiveCharacterTextSplitter
                ocr_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=300,
                    chunk_overlap=50
                )
                
                # 创建临时文档进行分块
                split_docs = ocr_splitter.split_documents([document])
                
                # 为每个分块添加ID
                import uuid
                for i, doc in enumerate(split_docs):
                    doc.metadata['chunk_id'] = str(uuid.uuid4())
                    doc.metadata['chunk_index'] = i
                
                logger.info(f"成功从图片提取文本，生成 {len(split_docs)} 个分块")
                return split_docs
                
        except pytesseract.TesseractError as e:
            logger.error(f"OCR处理失败: {file_path}, 错误: {str(e)}")
            return []
        except Exception as e:
            logger.error(f"处理图片失败: {file_path}, 错误: {str(e)}")
            return []
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """预处理图片以提高OCR质量"""
        preprocess_config = self.processing_config.get('image_preprocessing', {})
        
        # 转换为灰度图
        if preprocess_config.get('convert_to_grayscale', True):
            image = image.convert('L')
        
        # 调整对比度
        if preprocess_config.get('enhance_contrast', True):
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
        
        # 调整亮度
        if preprocess_config.get('enhance_brightness', False):
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(1.2)
        
        # 二值化
        if preprocess_config.get('binarize', False):
            threshold = preprocess_config.get('threshold', 128)
            image = image.point(lambda x: 255 if x > threshold else 0, '1')
        
        return image
    
    def _get_cache_key(self, file_path: str) -> str:
        """生成缓存键"""
        file_stat = os.stat(file_path)
        return f"{file_path}:{file_stat.st_mtime}:{file_stat.st_size}"
    
    def _validate_document(self, file_path: str, file_ext: str) -> bool:
        """验证文档有效性"""
        validation_config = self.processing_config.get('document_validation', {})
        if not validation_config.get('enabled', True):
            return True
        
        # 检查文件是否为有效的文本文件
        if file_ext in ['.txt', '.md', '.csv']:
            try:
                # 尝试以UTF-8编码打开文件，检查是否为有效文本
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(4096)  # 只读取一部分进行验证
            except UnicodeDecodeError:
                logger.warning(f"文件编码错误，不是有效文本文件: {file_path}")
                return False
        
        # PDF文件验证
        elif file_ext == '.pdf':
            try:
                # 简单验证PDF文件头
                with open(file_path, 'rb') as f:
                    header = f.read(4)
                    if header != b'%PDF':
                        logger.warning(f"无效的PDF文件: {file_path}")
                        return False
            except Exception as e:
                logger.warning(f"PDF文件验证失败: {file_path}, 错误: {str(e)}")
                return False
        
        # 文件大小验证
        max_size_mb = validation_config.get('max_size_mb', 100)
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > max_size_mb:
            logger.error(f"文件太大: {file_size_mb:.2f}MB，超过限制: {max_size_mb}MB")
            return False
        
        return True
    
    def _classify_document_type(self, document: Document) -> str:
        """根据文档内容分类文档类型"""
        try:
            content = document.page_content.lower()
            
            # 合同文档关键词
            if any(keyword in content for keyword in ['合同', '协议', '条款', '甲方', '乙方', '违约责任']):
                return 'contract'
            
            # 学术论文关键词
            if any(keyword in content for keyword in ['摘要', '关键词', '引言', '实验', '结论', '参考文献']):
                return 'academic_paper'
            
            # 技术文档关键词
            if any(keyword in content for keyword in ['教程', '指南', '使用说明', '配置', '安装', '功能']):
                return 'technical_doc'
            
            return 'default'
        except Exception:
            return 'default'
    
    def _split_documents(self, documents: List[Document], file_ext: str, doc_type: str) -> List[Document]:
        """根据文档类型选择合适的分块策略"""
        # 首先使用基于分隔符的分块
        if doc_type == 'contract':
            split_docs = self.contract_text_splitter.split_documents(documents)
        elif doc_type == 'academic_paper' or file_ext == '.pdf':
            split_docs = self.pdf_text_splitter.split_documents(documents)
        elif doc_type == 'technical_doc':
            split_docs = self.tech_text_splitter.split_documents(documents)
        else:
            split_docs = self.text_splitter.split_documents(documents)
        
        # 如果启用了语义分块，进行进一步优化
        if self._embedding_model and self.processing_config.get('semantic_chunking', {}).get('enabled', True):
            split_docs = self._perform_semantic_chunking(split_docs)
        
        return split_docs
    
    def _perform_semantic_chunking(self, split_docs: List[Document]) -> List[Document]:
        """基于语义相似度进行分块优化"""
        try:
            if len(split_docs) <= 1:
                return split_docs
            
            # 计算每个分块的嵌入
            chunks = [doc.page_content for doc in split_docs]
            embeddings = self._embedding_model.encode(chunks)
            
            # 计算相邻分块的相似度
            semantic_split_docs = []
            current_chunk = split_docs[0].page_content
            current_metadata = split_docs[0].metadata.copy()
            
            for i in range(1, len(split_docs)):
                similarity = np.dot(embeddings[i-1], embeddings[i]) / (
                    np.linalg.norm(embeddings[i-1]) * np.linalg.norm(embeddings[i])
                )
                
                # 如果相似度高且当前块大小适中，合并分块
                if similarity > 0.85 and len(current_chunk) < 2000:
                    current_chunk += "\n" + split_docs[i].page_content
                else:
                    # 添加当前块
                    new_doc = Document(page_content=current_chunk, metadata=current_metadata)
                    semantic_split_docs.append(new_doc)
                    # 开始新块
                    current_chunk = split_docs[i].page_content
                    current_metadata = split_docs[i].metadata.copy()
            
            # 添加最后一个块
            new_doc = Document(page_content=current_chunk, metadata=current_metadata)
            semantic_split_docs.append(new_doc)
            
            logger.info(f"语义分块完成，从 {len(split_docs)} 合并到 {len(semantic_split_docs)} 个分块")
            return semantic_split_docs
        except Exception as e:
            logger.warning(f"语义分块失败: {str(e)}，使用原始分块")
            return split_docs
    
    def _split_document_with_semantics(self, document: Document, doc_type: str) -> List[Document]:
        """使用语义分块方法分割文档，支持多种分块策略"""
        if not self._embedding_model:
            logger.warning("嵌入模型未初始化，无法使用语义分块")
            # 回退到字符分块
            return self._get_text_splitter_for_doc_type(doc_type).split_documents([document])
        
        # 获取文档类型的分块配置
        doc_type_config = self.processing_config.get('chunking_strategies', {}).get(doc_type, {})
        use_semantic = doc_type_config.get('use_semantic_chunking', False)
        strategy = doc_type_config.get('semantic_strategy', 'default')
        
        if not use_semantic:
            # 使用对应的字符分块器
            text_splitter = self._get_text_splitter_for_doc_type(doc_type)
            return text_splitter.split_documents([document])
        
        try:
            # 根据策略选择分块方法
            if strategy == 'paragraph_based':
                return self._semantic_chunking_by_paragraphs(document, doc_type)
            elif strategy == 'topic_modeling':
                return self._semantic_chunking_with_topic_modeling(document, doc_type)
            elif strategy == 'hierarchical':
                return self._hierarchical_semantic_chunking(document, doc_type)
            else:  # default
                return self._default_semantic_chunking(document, doc_type)
                
        except Exception as e:
            logger.error(f"语义分块失败: {str(e)}")
            # 回退到字符分块
            text_splitter = self._get_text_splitter_for_doc_type(doc_type)
            return text_splitter.split_documents([document])
    
    def _get_text_splitter_for_doc_type(self, doc_type: str):
        """根据文档类型获取对应的文本分块器"""
        if doc_type == 'contract':
            return self.contract_text_splitter
        elif doc_type == 'academic_paper':
            return self.pdf_text_splitter
        elif doc_type == 'technical_doc':
            return self.tech_text_splitter
        else:
            return self.text_splitter
    
    def _default_semantic_chunking(self, document: Document, doc_type: str) -> List[Document]:
        """默认的语义分块方法"""
        # 语义分块配置
        semantic_config = self.processing_config.get('semantic_chunking', {})
        max_chunk_size = semantic_config.get('max_chunk_size', 512)
        similarity_threshold = semantic_config.get('similarity_threshold', 0.3)
        min_chunk_size = semantic_config.get('min_chunk_size', 100)
        
        # 获取文本
        text = document.page_content
        # 基本的句子分割，优化中文分割
        sentences = self._split_text_into_sentences(text)
        
        # 过滤空句子
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            return [document]  # 如果没有句子，返回原始文档
        
        # 获取句子嵌入
        sentence_embeddings = self._embedding_model.encode(sentences, convert_to_tensor=True)
        
        # 计算句子之间的相似度并分块
        import numpy as np
        chunks = []
        current_chunk = [sentences[0]]
        
        for i in range(1, len(sentences)):
            # 计算当前句子与块中最后一个句子的相似度
            similarity = self._calculate_similarity(
                sentence_embeddings[i-1], sentence_embeddings[i]
            )
            
            # 检查当前块大小和相似度
            current_chunk_text = ' '.join(current_chunk)
            if (similarity < similarity_threshold or 
                len(current_chunk_text) + len(sentences[i]) > max_chunk_size):
                # 分割块
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentences[i]]
            else:
                current_chunk.append(sentences[i])
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        # 处理过小的块
        processed_chunks = self._merge_small_chunks(chunks, min_chunk_size)
        
        # 创建文档对象
        return self._create_chunk_documents(document, processed_chunks, 'default_semantic')
    
    def _semantic_chunking_by_paragraphs(self, document: Document, doc_type: str) -> List[Document]:
        """基于段落的语义分块，先按段落分割，再在必要时进行语义合并"""
        # 获取文本
        text = document.page_content
        
        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', text)  # 两个或多个换行符表示段落
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        if not paragraphs:
            return [document]
        
        # 获取段落嵌入
        paragraph_embeddings = self._embedding_model.encode(paragraphs, convert_to_tensor=True)
        
        # 配置
        semantic_config = self.processing_config.get('semantic_chunking', {})
        max_chunk_size = semantic_config.get('max_chunk_size', 1024)
        paragraph_similarity_threshold = semantic_config.get('paragraph_similarity_threshold', 0.6)
        
        # 合并段落以创建语义连贯的块
        import numpy as np
        chunks = []
        current_chunk = [paragraphs[0]]
        
        for i in range(1, len(paragraphs)):
            # 计算当前段落与块中最后一个段落的相似度
            similarity = self._calculate_similarity(
                paragraph_embeddings[i-1], paragraph_embeddings[i]
            )
            
            # 检查当前块大小
            current_chunk_text = '\n\n'.join(current_chunk)
            if (similarity < paragraph_similarity_threshold or 
                len(current_chunk_text) + len(paragraphs[i]) > max_chunk_size):
                # 分割块
                chunks.append(current_chunk_text)
                current_chunk = [paragraphs[i]]
            else:
                current_chunk.append(paragraphs[i])
        
        # 添加最后一个块
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        # 创建文档对象
        return self._create_chunk_documents(document, chunks, 'paragraph_semantic')
    
    def _semantic_chunking_with_topic_modeling(self, document: Document, doc_type: str) -> List[Document]:
        """使用简单的主题建模进行语义分块"""
        # 获取文本
        text = document.page_content
        
        # 分割句子
        sentences = self._split_text_into_sentences(text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            return [document]
        
        # 获取句子嵌入
        sentence_embeddings = self._embedding_model.encode(sentences, convert_to_tensor=True)
        
        # 使用简单的聚类来识别主题
        import numpy as np
        from sklearn.cluster import KMeans
        
        # 确定主题数量（简化版本）
        n_topics = min(5, max(2, len(sentences) // 10))
        
        try:
            # 应用K-means聚类
            kmeans = KMeans(n_clusters=n_topics, random_state=42, n_init=10)
            # 转换为numpy数组进行聚类
            if hasattr(sentence_embeddings, 'cpu'):
                embeddings_array = sentence_embeddings.cpu().numpy()
            else:
                embeddings_array = sentence_embeddings.numpy()
            
            clusters = kmeans.fit_predict(embeddings_array)
            
            # 根据聚类结果分组句子
            topic_groups = {}
            for i, cluster_id in enumerate(clusters):
                if cluster_id not in topic_groups:
                    topic_groups[cluster_id] = []
                topic_groups[cluster_id].append((i, sentences[i]))
            
            # 按句子顺序排序每个主题组
            for cluster_id in topic_groups:
                topic_groups[cluster_id].sort(key=lambda x: x[0])
            
            # 创建块
            chunks = []
            for cluster_id in sorted(topic_groups.keys()):
                sentences_in_topic = [sentence for _, sentence in topic_groups[cluster_id]]
                # 合并相同主题的句子
                chunk_text = ' '.join(sentences_in_topic)
                chunks.append(chunk_text)
            
            # 创建文档对象
            return self._create_chunk_documents(document, chunks, 'topic_modeling')
            
        except Exception as e:
            logger.warning(f"主题建模分块失败，回退到默认分块: {str(e)}")
            return self._default_semantic_chunking(document, doc_type)
    
    def _hierarchical_semantic_chunking(self, document: Document, doc_type: str) -> List[Document]:
        """层次化语义分块，先创建小块，再根据语义合并"""
        # 配置
        semantic_config = self.processing_config.get('semantic_chunking', {})
        base_chunk_size = semantic_config.get('base_chunk_size', 100)
        target_chunk_size = semantic_config.get('target_chunk_size', 512)
        merge_similarity_threshold = semantic_config.get('merge_similarity_threshold', 0.7)
        
        # 第一步：创建小块
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        base_splitter = RecursiveCharacterTextSplitter(
            chunk_size=base_chunk_size,
            chunk_overlap=20
        )
        small_chunks = base_splitter.split_text(document.page_content)
        
        if not small_chunks:
            return [document]
        
        # 获取小块嵌入
        chunk_embeddings = self._embedding_model.encode(small_chunks, convert_to_tensor=True)
        
        # 第二步：根据语义相似度合并小块
        import numpy as np
        merged_chunks = []
        current_chunk = [small_chunks[0]]
        
        for i in range(1, len(small_chunks)):
            # 计算当前块与最后一个块的相似度
            similarity = self._calculate_similarity(
                chunk_embeddings[i-1], chunk_embeddings[i]
            )
            
            # 检查当前块大小
            current_chunk_text = ' '.join(current_chunk)
            if (similarity < merge_similarity_threshold or 
                len(current_chunk_text) + len(small_chunks[i]) > target_chunk_size):
                # 完成当前块
                merged_chunks.append(current_chunk_text)
                current_chunk = [small_chunks[i]]
            else:
                current_chunk.append(small_chunks[i])
        
        # 添加最后一个块
        if current_chunk:
            merged_chunks.append(' '.join(current_chunk))
        
        # 创建文档对象
        return self._create_chunk_documents(document, merged_chunks, 'hierarchical')
    
    def _split_text_into_sentences(self, text: str) -> List[str]:
        """智能分割文本为句子，支持中英文"""
        # 中文句子分割正则表达式
        # 匹配句号、感叹号、问号后面跟着空格或换行符的情况
        pattern = r'(?<=[。！？.!?])\s*'
        sentences = re.split(pattern, text)
        
        # 处理可能的空句子和清理
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def _calculate_similarity(self, vec1, vec2) -> float:
        """计算两个向量之间的余弦相似度"""
        import numpy as np
        
        # 确保向量是numpy数组
        if hasattr(vec1, 'cpu'):
            vec1 = vec1.cpu().numpy()
        if hasattr(vec2, 'cpu'):
            vec2 = vec2.cpu().numpy()
        
        # 计算余弦相似度
        try:
            similarity = np.dot(vec1, vec2) / \
                        (np.linalg.norm(vec1) * np.linalg.norm(vec2) + 1e-10)
            return float(similarity)
        except Exception as e:
            logger.warning(f"相似度计算失败: {str(e)}")
            return 0.0
    
    def _merge_small_chunks(self, chunks: List[str], min_chunk_size: int) -> List[str]:
        """合并过小的块"""
        processed_chunks = []
        i = 0
        
        while i < len(chunks):
            chunk = chunks[i]
            # 如果块太小，尝试合并
            if len(chunk) < min_chunk_size and i < len(chunks) - 1:
                merged_chunk = chunk + ' ' + chunks[i+1]
                processed_chunks.append(merged_chunk)
                i += 2  # 跳过下一个块，因为已经合并了
            else:
                processed_chunks.append(chunk)
                i += 1
        
        return processed_chunks
    
    def _create_chunk_documents(self, original_doc: Document, chunks: List[str], chunk_method: str) -> List[Document]:
        """从文本块创建文档对象"""
        result_docs = []
        
        for i, chunk_text in enumerate(chunks):
            chunk_doc = Document(
                page_content=chunk_text,
                metadata={
                    **original_doc.metadata,
                    'chunk_id': str(uuid.uuid4()),
                    'chunk_index': i,
                    'chunk_method': chunk_method,
                    'total_chunks': len(chunks)
                }
            )
            result_docs.append(chunk_doc)
        
        logger.info(f"使用{chunk_method}方法将文档分割为 {len(result_docs)} 个块")
        return result_docs
    
    def _enhance_metadata(self, documents: List[Document], file_path: str, 
                         file_ext: str, doc_type: str) -> List[Document]:
        """增强文档元数据"""
        if not self.processing_config.get('metadata_extraction', {}).get('enabled', True):
            # 基本元数据
            for doc in documents:
                if 'source' not in doc.metadata:
                    doc.metadata['source'] = os.path.basename(file_path)
                if 'file_path' not in doc.metadata:
                    doc.metadata['file_path'] = file_path
                if 'file_type' not in doc.metadata:
                    doc.metadata['file_type'] = file_ext
                if 'chunk_id' not in doc.metadata:
                    doc.metadata['chunk_id'] = str(uuid.uuid4())
            return documents
        
        # 增强元数据提取
        enhanced_docs = []
        base_metadata = {
            'source': os.path.basename(file_path),
            'file_path': file_path,
            'file_type': file_ext,
            'doc_type': doc_type,
            'processed_at': time.strftime('%Y-%m-%d %H:%M:%S'),
            'language': self._detect_language(documents[0].page_content) if documents else 'unknown'
        }
        
        # 提取标题（从文件名或内容）
        base_metadata['title'] = self._extract_title(file_path, documents)

        # 解析 DOCX 核心属性（作者、创建/修改时间 等）
        try:
            ext_low = (file_ext or '').lower()
            if ('docx' in ext_low or ext_low == '.docx') and DocxDocument is not None:
                props = DocxDocument(file_path).core_properties
                author = getattr(props, 'author', None) or getattr(props, 'creator', None)
                created = getattr(props, 'created', None)
                modified = getattr(props, 'modified', None)
                last_modified_by = getattr(props, 'last_modified_by', None)
                title_prop = getattr(props, 'title', None)

                if author:
                    base_metadata['author'] = str(author)
                if created:
                    try:
                        base_metadata['created'] = created.isoformat()
                    except Exception:
                        base_metadata['created'] = str(created)
                if modified:
                    try:
                        base_metadata['modified'] = modified.isoformat()
                    except Exception:
                        base_metadata['modified'] = str(modified)
                if last_modified_by:
                    base_metadata['last_modified_by'] = str(last_modified_by)
                # 用文档属性标题补全（若前面未提取到）
                if title_prop and not base_metadata.get('title'):
                    base_metadata['title'] = str(title_prop)
        except Exception as e:
            logger.warning(f"DOCX 核心属性解析失败: {str(e)}")

        # 从文本内容中进一步提取作者/发布日期等（支持“Prepared by/Release Date”）
        try:
            sample_text = "\n".join([d.page_content for d in documents[:5]]) if documents else ""
            if sample_text:
                text_meta = extract_metadata_from_text(sample_text)
                # 映射作者
                if 'author' not in base_metadata and text_meta.get('author'):
                    base_metadata['author'] = text_meta['author']
                # 映射创建时间（优先Release Date/Release，然后Date）
                if 'created' not in base_metadata:
                    for k in ['created', 'release_date', 'release', 'date']:
                        if text_meta.get(k):
                            base_metadata['created'] = text_meta.get(k)
                            break
        except Exception as e:
            logger.warning(f"文本元数据提取失败: {str(e)}")
        
        # 为每个分块添加元数据
        for i, doc in enumerate(documents):
            metadata = {**doc.metadata, **base_metadata, 'chunk_index': i}
            
            # 提取分块特定信息
            metadata['chunk_id'] = str(uuid.uuid4())
            metadata['word_count'] = len(doc.page_content.split())
            metadata['char_count'] = len(doc.page_content)
            
            # 提取关键词（简单实现）
            # 将关键词列表转换为逗号分隔的字符串，因为向量存储只支持基本类型
            keywords_list = self._extract_keywords(doc.page_content)
            metadata['keywords'] = ', '.join(keywords_list) if keywords_list else ''
            
            # 提取章节信息（如果有）
            if doc_type == 'academic_paper':
                metadata['section'] = self._extract_section(doc.page_content)
        
        enhanced_docs.append(Document(page_content=doc.page_content, metadata=metadata))
        
        return enhanced_docs
    
    def _detect_language(self, text: str) -> str:
        """检测文本语言"""
        try:
            if len(text) < 20:
                return 'unknown'
            language = langdetect.detect(text)
            # 语言代码映射
            lang_map = {
                'zh-cn': 'zh', 'zh-tw': 'zh', 'zh': 'zh',
                'en': 'en', 'ja': 'ja', 'ko': 'ko'
            }
            return lang_map.get(language, language)
        except Exception:
            return 'unknown'
    
    def _extract_title(self, file_path: str, documents: List[Document]) -> str:
        """提取文档标题"""
        # 首先从文件名提取
        title = os.path.splitext(os.path.basename(file_path))[0]
        
        # 如果有内容，尝试从内容提取
        if documents and len(documents) > 0:
            first_lines = documents[0].page_content.split('\n')[:5]  # 取前5行
            for line in first_lines:
                line = line.strip()
                if len(line) > 5 and len(line) < 100 and (line.istitle() or line.isupper() or len(line) > 10):
                    title = line
                    break
        
        return title
    
    def _extract_keywords(self, text: str) -> List[str]:
        """简单关键词提取"""
        try:
            # 移除标点符号
            text = re.sub(r'[\W_]+', ' ', text)
            words = text.lower().split()
            
            # 过滤停用词（简单实现）
            stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这', 'the', 'a', 'an', 'and', 'or', 'but', 'is', 'are', 'was', 'were', 'in', 'on', 'at', 'to', 'for', 'of'}
            filtered_words = [word for word in words if word not in stop_words and len(word) > 1]
            
            # 计算词频
            from collections import Counter
            word_counts = Counter(filtered_words)
            
            # 返回前5个高频词
            return [word for word, count in word_counts.most_common(5)]
        except Exception:
            return []
    
    def _extract_section(self, text: str) -> str:
        """提取学术论文章节信息"""
        sections = {
            '摘要': 'abstract',
            '引言': 'introduction',
            '方法': 'methodology',
            '实验': 'experiment',
            '结果': 'results',
            '讨论': 'discussion',
            '结论': 'conclusion',
            '参考文献': 'references'
        }
        
        text_lower = text.lower()
        for keyword, section in sections.items():
            if keyword.lower() in text_lower:
                return section
        
        return 'other'
    
    def _extract_text_with_ocr(self, file_path: str) -> List[Document]:
        """使用多语言OCR提取文本"""
        try:
            logger.info(f"使用OCR处理文件: {file_path}")
            
            # 获取OCR配置
            ocr_config = self.processing_config.get('ocr', {})
            if not ocr_config.get('enabled', True):
                logger.warning("OCR功能已禁用")
                return []
            
            # 获取支持的语言
            languages = '+'.join(ocr_config.get('languages', ['chi_sim', 'eng']))
            
            if file_path.lower().endswith('.pdf'):
                # 处理PDF文件
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                all_text = ""
                
                for page_num in range(min(len(doc), 50)):  # 限制处理页数
                    page = doc[page_num]
                    # 先尝试提取文本
                    text = page.get_text()
                    if text.strip():
                        all_text += text + "\n"
                    else:
                        # 如果没有提取到文本，尝试OCR
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        # 使用预处理方法提高OCR质量
                        processed_img = self._preprocess_image(img)
                        ocr_text = pytesseract.image_to_string(processed_img, lang=languages)
                        all_text += ocr_text + "\n"
                
                doc.close()
            else:
                # 处理图片文件
                # 直接调用_process_image_file方法
                return self._process_image_file(file_path)
                
                # 如果_process_image_file返回空，备用方案
                img = Image.open(file_path)
                processed_img = self._preprocess_image(img)
                all_text = pytesseract.image_to_string(processed_img, lang=languages)
            
            if all_text.strip():
                # 创建文档并添加增强元数据
                metadata = {
                    'source': os.path.basename(file_path),
                    'file_path': file_path,
                    'file_type': os.path.splitext(file_path)[1].lower(),
                    'extraction_method': 'ocr',
                    'ocr_languages': languages,
                    'processed_at': time.strftime('%Y-%m-%d %H:%M:%S')
                }
                
                # 分块处理OCR提取的文本
                from langchain_text_splitters import RecursiveCharacterTextSplitter
                ocr_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=300,
                    chunk_overlap=50
                )
                
                # 创建临时文档进行分块
                temp_doc = Document(page_content=all_text, metadata=metadata)
                split_docs = ocr_splitter.split_documents([temp_doc])
                
                # 为每个分块添加ID
                for i, doc in enumerate(split_docs):
                    doc.metadata['chunk_id'] = str(uuid.uuid4())
                    doc.metadata['chunk_index'] = i
                
                logger.info(f"OCR成功提取文本，生成 {len(split_docs)} 个分块")
                return split_docs
            else:
                logger.warning(f"OCR未能提取到文本: {file_path}")
                return []
        except Exception as e:
            logger.error(f"OCR处理失败: {str(e)}")
            return []
    
    def load_directory(self, directory_path=None, progress_callback=None) -> List[Document]:
        """加载目录中的所有文档，支持并发处理和进度跟踪"""
        directory_path = directory_path or self.config.DOCUMENTS_PATH
        
        if not os.path.exists(directory_path):
            logger.warning(f"目录不存在: {directory_path}")
            if progress_callback:
                progress_callback(1.0, "目录不存在")
            return []
        
        # 获取所有文件
        all_files = []
        try:
            for root, _, files in os.walk(directory_path):
                # 过滤忽略的文件/目录
                if self._should_ignore(root):
                    continue
                
                for file in files:
                    # 过滤文件类型
                    if not self._should_process_file(file):
                        continue
                    
                    file_path = os.path.join(root, file)
                    all_files.append(file_path)
            
            total_files = len(all_files)
            logger.info(f"找到 {total_files} 个可处理的文件")
            
            # 如果没有文件，直接返回
            if total_files == 0:
                if progress_callback:
                    progress_callback(1.0, "没有找到可处理的文件")
                return []
            
            # 根据配置选择并发或顺序处理
            if self._executor and total_files > 3:  # 文件数量大于3时使用并发
                return self._process_files_concurrently(all_files, progress_callback)
            else:
                return self._process_files_sequentially(all_files, progress_callback)
                
        except Exception as e:
            logger.error(f"加载目录文档失败: {directory_path}, 错误: {str(e)}")
            if progress_callback:
                progress_callback(1.0, f"处理失败: {str(e)}")
            return []
    
    def _should_ignore(self, path: str) -> bool:
        """检查是否应该忽略该目录"""
        ignore_patterns = self.processing_config.get('ignore_patterns', [])
        ignore_dirs = ignore_patterns.get('directories', ['.git', '__pycache__', 'node_modules', '.venv'])
        
        for ignore_dir in ignore_dirs:
            if ignore_dir in path:
                return True
        return False
    
    def _should_process_file(self, filename: str) -> bool:
        """检查文件是否应该被处理"""
        # 获取文件扩展名
        file_ext = os.path.splitext(filename)[1].lower()
        
        # 检查忽略的文件类型
        ignore_patterns = self.processing_config.get('ignore_patterns', {})
        ignore_exts = ignore_patterns.get('extensions', ['.tmp', '.temp', '.swp', '.bak'])
        if file_ext in ignore_exts:
            return False
        
        # 检查忽略的文件名模式
        ignore_filenames = ignore_patterns.get('filenames', ['Thumbs.db', '.DS_Store'])
        if filename in ignore_filenames:
            return False
        
        # 检查允许的文件类型
        allowed_exts = self.processing_config.get('allowed_extensions', [
            '.txt', '.md', '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', 
            '.jpg', '.jpeg', '.png', '.gif'
        ])
        return file_ext in allowed_exts
    
    def _process_files_concurrently(self, files: List[str], progress_callback) -> List[Document]:
        """并发处理文件"""
        all_documents = []
        processed_count = 0
        total_files = len(files)
        
        try:
            with self._executor:
                # 提交所有任务
                future_to_file = {self._executor.submit(self.load_document, file): file for file in files}
                
                # 处理完成的任务
                for future in as_completed(future_to_file):
                    file = future_to_file[future]
                    try:
                        docs = future.result()
                        all_documents.extend(docs)
                        processed_count += 1
                        
                        # 更新进度
                        progress = processed_count / total_files
                        if progress_callback:
                            progress_callback(progress, f"处理中: {os.path.basename(file)} ({processed_count}/{total_files})")
                            
                    except Exception as e:
                        logger.error(f"处理文件失败 {file}: {str(e)}")
                        processed_count += 1
                        # 即使失败也要更新进度
                        if progress_callback:
                            progress = processed_count / total_files
                            progress_callback(progress, f"处理失败: {os.path.basename(file)} ({processed_count}/{total_files})")
            
            logger.info(f"并发处理完成，成功加载 {len(all_documents)} 个文档块")
            if progress_callback:
                progress_callback(1.0, f"处理完成，共加载 {len(all_documents)} 个文档块")
            
            return all_documents
        except Exception as e:
            logger.error(f"并发处理文件失败: {str(e)}")
            if progress_callback:
                progress_callback(1.0, f"并发处理失败: {str(e)}")
            return all_documents  # 返回已处理的文档
    
    def _process_files_sequentially(self, files: List[str], progress_callback) -> List[Document]:
        """顺序处理文件"""
        all_documents = []
        total_files = len(files)
        
        for i, file in enumerate(files):
            try:
                docs = self.load_document(file)
                all_documents.extend(docs)
                
                # 更新进度
                progress = (i + 1) / total_files
                if progress_callback:
                    progress_callback(progress, f"处理中: {os.path.basename(file)} ({i+1}/{total_files})")
                    
            except Exception as e:
                logger.error(f"处理文件失败 {file}: {str(e)}")
                # 即使失败也要更新进度
                if progress_callback:
                    progress = (i + 1) / total_files
                    progress_callback(progress, f"处理失败: {os.path.basename(file)} ({i+1}/{total_files})")
        
        logger.info(f"顺序处理完成，成功加载 {len(all_documents)} 个文档块")
        if progress_callback:
            progress_callback(1.0, f"处理完成，共加载 {len(all_documents)} 个文档块")
        
        return all_documents
    
    def _get_loader_cls(self, file_extension: str) -> Optional[Any]:
        """根据文件扩展名获取对应的加载器类"""
        loader_mapping = {
            ".txt": TextLoader,
            ".md": TextLoader,
            ".pdf": PDFLoader,
            ".docx": UnstructuredWordDocumentLoader,
            ".doc": UnstructuredWordDocumentLoader,
            ".xlsx": UnstructuredExcelLoader,
            ".xls": UnstructuredExcelLoader,
            ".csv": CSVLoader
        }
        
        loader_cls = loader_mapping.get(file_extension.lower())
        if not loader_cls:
            logger.warning(f"不支持的文件类型: {file_extension}")
            return None
        
        return loader_cls
    
    def split_documents(self, documents):
        """将文档分割成小块"""
        if not documents:
            return []
        
        try:
            # 为PDF文档使用特殊的分割策略
            split_docs = []
            normal_docs = []
            pdf_docs = []
            pdf_split_docs = []  # 初始化变量，确保无论是否有PDF文档都已定义
            
            # 分离PDF文档和其他文档
            for doc in documents:
                if hasattr(doc, 'metadata') and 'source' in doc.metadata and doc.metadata['source'].lower().endswith('.pdf'):
                    pdf_docs.append(doc)
                else:
                    normal_docs.append(doc)
            
            # 先分割非PDF文档
            if normal_docs:
                normal_split_docs = self.text_splitter.split_documents(normal_docs)
                split_docs.extend(normal_split_docs)
                logger.info(f"成功分割非PDF文档，从 {len(normal_docs)} 个原始文档分割为 {len(normal_split_docs)} 个小块")
            
            # 对PDF文档使用更稳健的分割策略
            if pdf_docs:
                logger.info(f"正在使用特殊策略分割 {len(pdf_docs)} 个PDF文档")
                
                pdf_split_docs = []
                for pdf_doc in pdf_docs:
                    try:
                        # 使用我们预定义的PDF专用分割器
                        doc_splits = self.pdf_text_splitter.split_documents([pdf_doc])
                        pdf_split_docs.extend(doc_splits)
                        logger.info(f"成功分割PDF文档，获得 {len(doc_splits)} 个小块")
                    except Exception as e:
                        logger.error(f"分割PDF文档失败: {str(e)}")
                        try:
                            # 尝试使用常规文本分割器
                            doc_splits = self.text_splitter.split_documents([pdf_doc])
                            pdf_split_docs.extend(doc_splits)
                            logger.warning(f"使用常规分割器成功分割PDF文档，获得 {len(doc_splits)} 个小块")
                        except Exception as inner_e:
                            logger.error(f"PDF文档备用分割方法也失败: {str(inner_e)}")
                            # 最后尝试：直接将整个文档作为一个块
                            if len(pdf_doc.page_content) > 0:
                                pdf_split_docs.append(pdf_doc)
                                logger.warning(f"对文档使用原始内容作为单个块")
            
            split_docs.extend(pdf_split_docs)
            
            # 为每个文档块添加分类信息
            classified_docs = []
            for doc in split_docs:
                try:
                    if hasattr(doc, 'metadata') and 'source' in doc.metadata:
                        file_path = doc.metadata['source']
                        # 使用文档分类器对文档进行分类
                        classification_result = document_classifier.classify_document(file_path, doc.page_content)
                        
                        # 合并原始元数据和分类信息
                        enhanced_metadata = {
                            **doc.metadata,
                            "document_type": classification_result.get("document_type", "unknown"),
                            "document_type_name": classification_result.get("document_type_name", "未知文档类型"),
                            "classification_confidence": classification_result.get("confidence", 0.0),
                            "classification_method": classification_result.get("classification_method", "unknown"),
                            "source_file": file_path,
                            "source_file_name": os.path.basename(file_path)
                        }
                        
                        # 添加分类详情（如果有）
                        if "details" in classification_result:
                            enhanced_metadata["classification_details"] = classification_result["details"]
                        
                        # 添加从分类器获取的文档元数据
                        if "metadata" in classification_result:
                            for key, value in classification_result["metadata"].items():
                                # 避免覆盖已有键
                                if key not in enhanced_metadata:
                                    enhanced_metadata[key] = value
                        
                        # 创建带有增强元数据的新文档
                        from langchain_core.documents import Document
                        classified_doc = Document(
                            page_content=doc.page_content,
                            metadata=enhanced_metadata
                        )
                        classified_docs.append(classified_doc)
                    else:
                        classified_docs.append(doc)
                except Exception as e:
                    logger.error(f"文档分类失败: {str(e)}")
                    classified_docs.append(doc)  # 如果分类失败，使用原始文档
            
            # 关键修复：确保即使分割后得到0个文档块，也至少返回原始文档
            if not classified_docs:
                logger.warning(f"文档分割后得到0个文档块，返回原始文档")
                return documents
            
            logger.info(f"成功分割并分类所有文档，从 {len(documents)} 个原始文档分割为 {len(classified_docs)} 个小块")
            return classified_docs
        except Exception as e:
            logger.error(f"分割文档失败: {str(e)}")
            # 作为最后的后备方案，返回原始文档
            return documents

# 创建文档加载器实例
document_loader = DocumentLoader()